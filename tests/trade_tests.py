import warnings
warnings.filterwarnings("ignore")
import pandas as pd
import numpy as np
from sklearn.model_selection import train_test_split, learning_curve
from sklearn.linear_model import LinearRegression
from sklearn.tree import DecisionTreeRegressor
from sklearn.ensemble import RandomForestRegressor, ExtraTreesRegressor
from sklearn.metrics import mean_squared_error, r2_score
from sklearn.preprocessing import PolynomialFeatures, OneHotEncoder
from sklearn.compose import ColumnTransformer
import plotly.graph_objs as go
import plotly.io as pio
from itertools import combinations
from metaapi_cloud_sdk import MetaApi
import asyncio
import joblib
import sklearn
import os
from docx import Document
from sklearn.model_selection import cross_val_score


from datetime import datetime, timedelta
token = os.getenv('TOKEN') or 'eyJhbGciOiJSUzUxMiIsInR5cCI6IkpXVCJ9.eyJfaWQiOiI2YjI0NTQ0ZWYzMWI0NzQ4NWMxNzQ1NmUzNzdmYTlhZiIsInBlcm1pc3Npb25zIjpbXSwiYWNjZXNzUnVsZXMiOlt7ImlkIjoidHJhZGluZy1hY2NvdW50LW1hbmFnZW1lbnQtYXBpIiwibWV0aG9kcyI6WyJ0cmFkaW5nLWFjY291bnQtbWFuYWdlbWVudC1hcGk6cmVzdDpwdWJsaWM6KjoqIl0sInJvbGVzIjpbInJlYWRlciIsIndyaXRlciJdLCJyZXNvdXJjZXMiOlsiKjokVVNFUl9JRCQ6KiJdfSx7ImlkIjoibWV0YWFwaS1yZXN0LWFwaSIsIm1ldGhvZHMiOlsibWV0YWFwaS1hcGk6cmVzdDpwdWJsaWM6KjoqIl0sInJvbGVzIjpbInJlYWRlciIsIndyaXRlciJdLCJyZXNvdXJjZXMiOlsiKjokVVNFUl9JRCQ6KiJdfSx7ImlkIjoibWV0YWFwaS1ycGMtYXBpIiwibWV0aG9kcyI6WyJtZXRhYXBpLWFwaTp3czpwdWJsaWM6KjoqIl0sInJvbGVzIjpbInJlYWRlciIsIndyaXRlciJdLCJyZXNvdXJjZXMiOlsiKjokVVNFUl9JRCQ6KiJdfSx7ImlkIjoibWV0YWFwaS1yZWFsLXRpbWUtc3RyZWFtaW5nLWFwaSIsIm1ldGhvZHMiOlsibWV0YWFwaS1hcGk6d3M6cHVibGljOio6KiJdLCJyb2xlcyI6WyJyZWFkZXIiLCJ3cml0ZXIiXSwicmVzb3VyY2VzIjpbIio6JFVTRVJfSUQkOioiXX0seyJpZCI6Im1ldGFzdGF0cy1hcGkiLCJtZXRob2RzIjpbIm1ldGFzdGF0cy1hcGk6cmVzdDpwdWJsaWM6KjoqIl0sInJvbGVzIjpbInJlYWRlciIsIndyaXRlciJdLCJyZXNvdXJjZXMiOlsiKjokVVNFUl9JRCQ6KiJdfSx7ImlkIjoicmlzay1tYW5hZ2VtZW50LWFwaSIsIm1ldGhvZHMiOlsicmlzay1tYW5hZ2VtZW50LWFwaTpyZXN0OnB1YmxpYzoqOioiXSwicm9sZXMiOlsicmVhZGVyIiwid3JpdGVyIl0sInJlc291cmNlcyI6WyIqOiRVU0VSX0lEJDoqIl19LHsiaWQiOiJjb3B5ZmFjdG9yeS1hcGkiLCJtZXRob2RzIjpbImNvcHlmYWN0b3J5LWFwaTpyZXN0OnB1YmxpYzoqOioiXSwicm9sZXMiOlsicmVhZGVyIiwid3JpdGVyIl0sInJlc291cmNlcyI6WyIqOiRVU0VSX0lEJDoqIl19LHsiaWQiOiJtdC1tYW5hZ2VyLWFwaSIsIm1ldGhvZHMiOlsibXQtbWFuYWdlci1hcGk6cmVzdDpkZWFsaW5nOio6KiIsIm10LW1hbmFnZXItYXBpOnJlc3Q6cHVibGljOio6KiJdLCJyb2xlcyI6WyJyZWFkZXIiLCJ3cml0ZXIiXSwicmVzb3VyY2VzIjpbIio6JFVTRVJfSUQkOioiXX0seyJpZCI6ImJpbGxpbmctYXBpIiwibWV0aG9kcyI6WyJiaWxsaW5nLWFwaTpyZXN0OnB1YmxpYzoqOioiXSwicm9sZXMiOlsicmVhZGVyIl0sInJlc291cmNlcyI6WyIqOiRVU0VSX0lEJDoqIl19XSwidG9rZW5JZCI6IjIwMjEwMjEzIiwiaW1wZXJzb25hdGVkIjpmYWxzZSwicmVhbFVzZXJJZCI6IjZiMjQ1NDRlZjMxYjQ3NDg1YzE3NDU2ZTM3N2ZhOWFmIiwiaWF0IjoxNzI2NDY4MzM2LCJleHAiOjE3MzQyNDQzMzZ9.j_Zg0Uy4ECsdDMg1lj5nbPX8yub5Oq00sUwNt0H1f5wpvDGmhKeTY2IHDmqY52AECkqQkr-vRPxLsewezVTcTuXY88noJrZPxkdD5q5X9cvBTtwt3tgeisXZn3nGqvoZTXX4mLWsuGiRzWFIyFDLLR2KemTNN5fM0YYjIxqB5UgBAUQHZQqNu6CcO-SYOGUMdUHdpMacFkN3u-AApEjjldo4GQQNeCBzX-DIVb6k0PUtybQ_MTH9rdEf-0PqiKqlHxWtckO073uPb5pEXaObV8BWpHypPlI4bobEvVXhjnWBv69l_RfBpG2EJvHZ_5kVOjDGutHFifPliCtyvNLXiI2EY9Wia41XX4O_KMBXbTg_KqrEEvJeCZc-R_Fpjq-iEGQOkPjY8eiLg1FW0au21Hj0JzSjuzV__8kiQ5hGQNmbMiwPnefn3OHmn7M-cHfHF0wjCm5nIQo2Q0ZEjjL4ltvhArrkln5-3QsjT-Oa4g9bhdDka15EOxt_hmAhY7v9cjGySL4i5Ax76avPd8DJglIgQfsomaST65N68Ao4XnJHQKQe7qhMyAjaoMLRDp7k39cO2S_Yq4R7pv9C62qwfYeUvc3qf8Atr9OXDPV-8r4GMaoSj3Fczbp9yY9XtoEMZtIc92wCXxOlzntKGvpTQhnF3vlpp-rDdweS6nyA-5c'
#accountId = os.getenv('ACCOUNT_ID') or 'df662a60-74f5-4f40-a356-622e3f20c88d'
accountId = os.getenv('ACCOUNT_ID') or '52273464-c640-4cf1-a638-ec1da0a508a9'


symbol_list = [
    'XAUUSDm',  # Gold/US Dollar (Commodity)
    'XAGUSDm',
    'GBPAUDm' ,  # British Pound/Australian Dollar (Minor)
    'BTCUSDm',
    'EURUSDm',  # Euro/US Dollar (Major)
    'GBPUSDm', 
]
data = {}
timeframe='15m'
pages=8
n_estimators=1
min_samples_leaf=1
shuffle=True
max_depth=50
test_size=0.01

def decimal_places(number):
    # Convert the number to a string
    num_str = str(number)
    
    # Check if there is a decimal point
    if '.' in num_str:
        # Find the index of the decimal point
        decimal_index = num_str.index('.')
        
        # Count the characters after the decimal point
        num_decimal_places = len(num_str) - decimal_index - 1
        
        return num_decimal_places
    else:
        # If there is no decimal point, return 0
        return 0


# Function to add nested dictionary content to a .docx file
def add_nested_dict_to_docx(info_dict, filename=f'output{pages}{timeframe}.docx'):
    doc = Document()
    doc.add_heading('AI analysis and market forecast For timeframe', 0)
    for symbol, details in info_dict.items():
        doc.add_heading(symbol, level=1)
        for key, value in details.items():
            doc.add_paragraph(f"{key}: {value}", style='List Number')
    doc.save(filename)
    print(f"Information added to {filename} successfully.")

# Function to add data to the dictionary
def add_symbol_data(symbol, r2_2, mse_2, r2_2_1, mse_2_1, r2_2_2, mse_2_2, next_close, next_low, next_high, actual_close, actual_low, actual_high, diff2, diff1):
    if symbol not in data:
        data[symbol] = {}
    
    data[symbol] = {
        'r2_2': r2_2,
        'mse_2': mse_2,
        'r2_2_1': r2_2_1,
        'mse_2_1': mse_2_1,
        'r2_2_2': r2_2_2,
        'mse_2_2': mse_2_2,
        'next_close': next_close,
        'next_low': next_low,
        'next_high': next_high,
        'actual': {
            'close': actual_close,
            'low': actual_low,
            'high': actual_high
        },
        'diff2': diff2,
        'diff1': diff1
    }

def prepare(df):
    #df=add_stop_losse(df)
    df=df.drop(columns=['symbol','timeframe','brokerTime'])

    df['time'] = pd.to_datetime(df['time'])
    df['time'] = df['time'].astype(int)// 10**99 
    # Generate new features and handle data preparation
    df_new =df

    return df_new
async def main2(timeframe,pages):
    print('Up and runing')
    api = MetaApi(token)
    account = await api.metatrader_account_api.get_account(accountId)
    initial_state = account.state
    deployed_states = ['DEPLOYING', 'DEPLOYED']

    if initial_state not in deployed_states:
        # wait until account is deployed and connected to broker
        print('Deploying account')
        await account.deploy()
        print('Waiting for API server to connect to broker (may take a few minutes)')
        await account.wait_connected()
    # Connect to MetaApi API
    connection = account.get_rpc_connection()
    await connection.connect()

    # Wait until terminal state synchronized to the local state
    print('Waiting for SDK to synchronize to terminal state (may take some time depending on your history size)')
    await connection.wait_synchronized()
    
    for symbol in symbol_list:
        print(symbol)
        trades =await connection.get_positions()#connection.get_orders()
        if len(trades)>=12:
            print(f'There are more than 10 runing trades, Total is :{len(trades)}')

        else:
            #sentiment_results = float(cal_sentiment_textblob(symbol))
            #if sentiment_results>0.25 or sentiment_results<0.25:
            try:
                try:
                    # Fetch historical price data
                    candles = await account.get_historical_candles(symbol=symbol, timeframe=timeframe, start_time=None, limit=3)

                    print('Fetched the latest candle data successfully')
                except Exception as e:
                    raise e
                try:
                    if not isinstance(candles, str):
                        df=pd.DataFrame(candles)
                    else:
                        
                        df=pd.DataFrame()
                except Exception as e:
                    raise e

                if not df.empty:
                    # Drop unwanted columns
                    df = df.drop(columns=['symbol', 'timeframe', 'brokerTime'])

                    # Convert 'time' to datetime and create a new time feature
                    df['time'] = pd.to_datetime(df['time'])
                    df['time2'] = df['time'].astype(int) // 10**99  # Create a numerical time feature

                    #df['SMA_10'] = df['close'].rolling(window=10).mean()
                    #df['SMA_50'] = df['close'].rolling(window=50).mean()
    
                    # Set 'time' as the index
                    df.set_index('time', inplace=True)
                    predictions_features=np.array(df.iloc[-1]).reshape(1, -1)

                    model_close,_= joblib.load(f'../Regressors/Close/model{symbol}{timeframe}close.pkl')
                    model_high,_= joblib.load(f'../Regressors/High/model{symbol}{timeframe}high.pkl')
                    model_low,_= joblib.load(f'../Regressors/Low/model{symbol}{timeframe}low.pkl')

                    next_close=model_close.predict(predictions_features)[0].round(decimal_places(df['close'].iloc[-2]))
                    next_low=model_low.predict(predictions_features)[0].round(decimal_places(df['close'].iloc[-2]))
                    next_high=model_high.predict(predictions_features)[0].round(decimal_places(df['close'].iloc[-2]))
                    print(next_close,next_low,next_high)
                    previous_close= df['close'].iloc[-1]
                    print(f'Previous close : {previous_close}')
                    symbol_list_lag= [
                        {'BTCUSDm': 100},
                        {'GBPUSDm': 0.00048},
                        {'EURUSDm': 0.00044},
                        {'AUDUSDm': 0.00048},
                        {'XAUUSDm':1},
                        {'GBPAUDm':0.00181},
                        {'XAGUSDm':0.050},
                    ]

                    for item in symbol_list_lag:
                        if symbol in item:
                            lag_size= item[symbol]
                        else:
                            lag_size=0.0004

                    
                    if next_close>previous_close and next_close<next_high and next_close>next_low and (next_close-previous_close)>lag_size and ((next_close-previous_close)<(lag_size*4)):
                        stop_loss=None#next_low-(lag_size*4)
                        #take_profit=trademax-(lag_size/2)
                        try:
                            
                            result = await connection.create_market_buy_order(
                                symbol=symbol,
                                volume=0.01,
                                stop_loss=stop_loss,
                                take_profit=next_close,
                            )
                            print(f'Buy_Signal (T)   :Buy Trade successful For Symbol :{symbol}')
                            
                            Trader_success=True
                        except Exception as err:
                            print('Trade failed with error:')
                            print(api.format_error(err))
                    if next_close<previous_close and next_close<next_high and next_close>next_low and (previous_close-next_close)>lag_size  and ((previous_close-next_close)<(lag_size*4)):
                        stop_loss=None#(next_high+lag_size*4)
                        #take_profit=trademax+(lag_size/2)
                        try:
                            
                            result = await connection.create_market_sell_order(
                                symbol=symbol,
                                volume=0.01,
                                stop_loss=stop_loss,
                                take_profit=next_close,
                            )
                            print(f'Sell Signal (T)   :Sell Trade successful For Symbol :{symbol}')
                            Trader_success=True

                        except Exception as err:
                            #raise err
                            print('Trade failed with error:')
                            print(api.format_error(err))
                    else:
                        print('No trade conditions passed, so no trade placed')
                print('*'*20)
                print('*'*20)
                print('*'*20)
            except Exception as e:
                print(f"An error occurred: {e}")
#def main():
asyncio.run(main2(timeframe,pages))